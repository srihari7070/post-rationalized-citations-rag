"""
Assemble the markdown chapter files into a single Word document.

Markdown is the versioned source, because it diffs cleanly in git. The .docx is a build
artefact for reading and for pasting into the official template.

    python document/build_draft.py
    python document/build_draft.py --version v1
"""
import argparse
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

DOC_DIR = Path(__file__).parent
DRAFTS = DOC_DIR / "drafts"

# Chapters in reading order. Files absent from a version are skipped.
CHAPTERS = [
    ("01_introduction.md", "1 Introduction"),
    ("02_literature.md", "2 Literature Review"),
    ("03_methodology.md", "3 Methodology"),
    ("04_findings.md", "4 Findings"),
    ("05_discussion.md", "5 Discussion"),
    ("06_conclusion.md", "6 Conclusion"),
]


def add_table(doc, rows):
    """Render a markdown table. First row is the header, second is the separator."""
    header = [c.strip() for c in rows[0].strip("|").split("|")]
    body = [[c.strip() for c in r.strip("|").split("|")] for r in rows[2:]]

    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(re.sub(r"\*\*(.+?)\*\*", r"\1", h))
        run.bold = True
        run.font.size = Pt(9)
    for row in body:
        cells = t.add_row().cells
        for i, val in enumerate(row[:len(header)]):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(
                re.sub(r"\*\*(.+?)\*\*", r"\1", val))
            run.font.size = Pt(9)
    doc.add_paragraph()


def add_rich(par, text):
    """Bold spans marked with ** **, everything else plain."""
    for i, part in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if part:
            par.add_run(part).bold = bool(i % 2)


def render_markdown(doc, md):
    lines = md.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line:
            i += 1
            continue

        # tables
        if line.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            if len(block) >= 3:
                add_table(doc, block)
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(36)
            r = p.add_run(line[2:])
            r.italic = True
        elif re.match(r"^[-*] ", line):
            p = doc.add_paragraph(style="List Bullet")
            add_rich(p, line[2:])
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_rich(p, re.sub(r"^\d+\. ", "", line))
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            add_rich(p, line)
        i += 1


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--version", default="v1")
    args = ap.parse_args()

    src = DRAFTS / args.version
    if not src.exists():
        print(f"No draft directory at {src}")
        return

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(10)
    style.paragraph_format.line_spacing = 1.15

    # Title block
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Detecting and Correcting Post-Rationalised Citations\n"
                  "in Retrieval-Augmented Generation")
    r.bold = True
    r.font.size = Pt(18)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run(f"Working draft {args.version}  ·  {date.today().isoformat()}")
    sr.font.size = Pt(10)
    sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    n = doc.add_paragraph()
    n.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = n.add_run("Draft for revision. Not for submission in this form.")
    nr.italic = True
    nr.font.size = Pt(9)
    nr.font.color.rgb = RGBColor(0x99, 0x66, 0x00)

    doc.add_page_break()

    included, missing = [], []
    for fname, label in CHAPTERS:
        path = src / fname
        if not path.exists():
            missing.append(label)
            continue
        render_markdown(doc, path.read_text())
        doc.add_page_break()
        included.append(label)

    if missing:
        doc.add_heading("Not yet drafted", level=1)
        for label in missing:
            doc.add_paragraph(label, style="List Bullet")

    out = src / f"thesis_draft_{args.version}_{date.today().isoformat()}.docx"
    doc.save(out)

    words = sum(len((src / f).read_text().split())
                for f, _ in CHAPTERS if (src / f).exists())
    print(f"included: {', '.join(included)}")
    if missing:
        print(f"missing : {', '.join(missing)}")
    print(f"words   : {words:,}")
    print(f"saved   : {out}")


if __name__ == "__main__":
    main()
