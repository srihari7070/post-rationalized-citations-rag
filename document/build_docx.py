"""
Build the thesis as a Word document using the SRH template.

Clones `Thesis Report Template MSc New.docx` so that page setup, margins, headers,
footers and the built-in Title / Heading / Table Grid styles all carry over, then
clears the template's instructional body and writes the real chapters in.

Chapters are read from document/dictation/dictated where the author's own dictated
text exists (the real thesis prose), falling back to document/drafts/v2 and then
document/drafts/v1 for any chapter not yet dictated.

    python document/build_docx.py
"""
import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
TEMPLATE = ROOT / "Thesis Report Template MSc New.docx"
OUT = ROOT / "Thesis_Srihari_Ananthan.docx"
DICTATED, V1, V2 = ROOT / "dictation/dictated", ROOT / "drafts/v1", ROOT / "drafts/v2"

CHAPTERS = ["01_introduction", "02_literature", "03_methodology",
            "04_findings", "05_discussion", "06_conclusion"]

TITLE = "Correcting Post-Rationalised Citations in Retrieval-Augmented Generation"
SHORT_TITLE = "Correcting Post-Rationalised Citations in RAG"  # for the running header only, so it fits on one line
SUBTITLE = "A causal feedback loop and the generator property it makes visible"
AUTHOR = "Srihari Ananthan"
MATRIC = "100001648"
SUPERVISOR = "Prof. Dr. Joel Dokmegang"
ASSOCIATE = "Maximilian Erdmann Sanchez, Startup Insider GmbH"
PROGRAMME = "Master of Science, Computer Science, Big Data and Artificial Intelligence"


def no_auto_number(par):
    """The template's Heading 1/2/3 styles carry Word's own automatic outline
    numbering (numPr) linked to a multilevel list. Every chapter/section number
    in this document is already typed as real text (from the source markdown's
    own '2.5 Attempts to correct citations' headers), so leaving the style's
    numbering active double-numbers every heading. numId 0 disables list
    numbering for this one paragraph without touching the style itself."""
    pPr = par._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId"); numId.set(qn("w:val"), "0")
    numPr.append(ilvl); numPr.append(numId)
    pPr.append(numPr)


def add_heading(doc, text, style):
    p = doc.add_paragraph(text, style=style)
    no_auto_number(p)
    return p


def clear_body(doc):
    """Remove every block in the body but keep sectPr, which holds page setup."""
    body = doc.element.body
    for child in list(body):
        if not child.tag.endswith("}sectPr"):
            body.remove(child)


RED = RGBColor(0xC0, 0x00, 0x00)


def add_runs(par, text):
    """Inline **bold**, *italic*, `code`, [^N] footnote markers (rendered as a
    visual superscript number -- python-docx cannot create a real Word footnote
    object; see refs/README.md for the required manual Alt+Ctrl+F pass), and
    {{...}} red-highlight markers for text flagged for the author's own review
    (e.g. newly added citations, not yet confirmed to fit). Markdown links
    reduce to their label."""
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    for piece in re.split(r"(\{\{.*?\}\}|\*\*[^*]+\*\*|(?<!\*)\*[^*]+\*(?!\*)|`[^`]+`|\[\^\d+\])", text):
        if not piece:
            continue
        if piece.startswith("{{") and piece.endswith("}}"):
            start = len(par.runs)
            add_runs(par, piece[2:-2])
            for r in par.runs[start:]:
                r.font.color.rgb = RED
        elif piece.startswith("**") and piece.endswith("**"):
            par.add_run(piece[2:-2]).bold = True
        elif piece.startswith("*") and piece.endswith("*"):
            par.add_run(piece[1:-1]).italic = True
        elif piece.startswith("`") and piece.endswith("`"):
            r = par.add_run(piece[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9)
        elif re.match(r"^\[\^\d+\]$", piece):
            r = par.add_run(piece[2:-1])
            r.font.superscript = True
        else:
            par.add_run(piece)


TABLE_CAP_RE = re.compile(r"^(Table\s+(?:\d+|[A-Z])\.\d+[a-z]?)\b(.*)$")
FIG_CAP_RE = re.compile(r"^!\[(.*?)\]\(")
FORMULA_RE = re.compile(r"^FORMULA\s+(\d+\.\d+):\s*(.+)$")


def add_formula(doc, number, expr):
    """python-docx cannot create real Word equation objects, so a formula is a
    small centered, numbered text paragraph instead of an image -- consistent
    with the template's own allowance for typed formulas at first definition."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    label = p.add_run(f"Formula {number}   ")
    label.bold = True
    expr_run = p.add_run(expr)
    expr_run.font.name = "Consolas"
    expr_run.font.size = Pt(11)
    return p


def style_as_table_caption(par):
    """If par's text is a 'Table X.Y ...' announcement sentence, turn it into a
    real centred Caption-styled label (bold number, per the template's own
    figure-caption rule at line 336) instead of leaving it as flowing body text."""
    m = TABLE_CAP_RE.match(par.text.strip())
    if not m:
        return None
    label, rest = m.groups()
    for r in list(par.runs):
        r._element.getparent().remove(r._element)
    par.style = "Caption"
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run(label).bold = True
    if rest:
        par.add_run(rest)
    return label + rest


def add_table(doc, rows):
    header, body = rows[0], rows[2:]          # rows[1] is the |---| separator
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    for i, cell in enumerate(header):
        p = t.rows[0].cells[i].paragraphs[0]
        add_runs(p, cell)
        for r in p.runs:
            r.bold = True
    for row in body:
        cells = t.add_row().cells
        for i, cell in enumerate(row[:len(header)]):
            add_runs(cells[i].paragraphs[0], cell)
    doc.add_paragraph()


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def write_markdown(doc, md):
    lines = md.split("\n")
    i, para = 0, []

    def flush():
        if para:
            add_runs(doc.add_paragraph(style="Normal"), " ".join(para))
            para.clear()

    while i < len(lines):
        line = lines[i].rstrip()

        # scaffolding notes for the author, not thesis content -- never render
        if line.strip() == "*Convert each to a Word footnote (Alt+Ctrl+F), Times New Roman 10pt, single-spaced.*":
            i += 1; continue
        if line.strip() == "---":
            i += 1; continue

        formula = FORMULA_RE.match(line.strip())
        if formula:
            flush()
            add_formula(doc, formula.group(1), formula.group(2))
            i += 1; continue

        img = re.match(r"^!\[(.*?)\]\(([^)]+)\)$", line.strip())
        if img:
            flush()
            caption, relpath = img.group(1), img.group(2)
            path = (ROOT / relpath).resolve()
            if path.exists():
                doc.add_picture(str(path), width=Inches(6.0))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = doc.add_paragraph(style="Caption")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_runs(cap, caption)
            else:
                add_runs(doc.add_paragraph(style="Normal"), f"[missing figure: {relpath}]")
            i += 1; continue

        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[i+1].strip()):
            flush()
            if doc.paragraphs:
                style_as_table_caption(doc.paragraphs[-1])
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i])); i += 1
            add_table(doc, rows); continue

        if line.startswith("#"):
            flush()
            level = len(line) - len(line.lstrip("#"))
            text = line.lstrip("#").strip()
            if level == 1:
                text = re.sub(r"^\d+\s+", "", text)
                small = {"and", "of", "the", "in", "on", "for", "an", "to", "as"}
                acronyms = {"rag", "prr", "gcr", "ccr", "llm", "ai", "api", "rq1", "rq2",
                            "rq3", "auroc", "cue-r", "ieee", "acm"}
                def cased(w, i):
                    core = re.sub(r"[,:;]+$", "", w)
                    if core.lower() in acronyms:
                        return core.upper() + w[len(core):]
                    return w.capitalize() if i == 0 or w.lower() not in small else w.lower()
                words = text.split()
                text = " ".join(cased(w, i) for i, w in enumerate(words))
                add_heading(doc, text, "Heading 1")
            else:
                add_heading(doc, text, f"Heading {min(level, 3)}")
            i += 1; continue

        if line.startswith(">"):
            flush()
            add_runs(doc.add_paragraph(style="Quote"), line.lstrip("> ").strip())
            i += 1; continue

        if re.match(r"^[-*]\s+", line):
            flush()
            add_runs(doc.add_paragraph(style="List Paragraph"), re.sub(r"^[-*]\s+", "", line))
            i += 1; continue

        if re.match(r"^\d+\.\s+", line):
            flush()
            add_runs(doc.add_paragraph(style="List Paragraph"), line.strip())
            i += 1; continue

        if not line.strip():
            flush(); i += 1; continue

        para.append(line.strip()); i += 1
    flush()


def collect_captions(texts):
    """Pre-scan the raw markdown (in the same order it will be written) for figure
    and table captions, so the List of Figures / List of Tables front-matter
    section can be built before the chapters that contain them. A table caption
    only counts when it immediately precedes a real table -- a stray in-prose
    back-reference like 'Table 4.3a already shows...' must not be listed."""
    figures, tables = [], []
    for text in texts:
        lines = text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            fm = FIG_CAP_RE.match(line)
            if fm:
                num, _, cap = fm.group(1).partition(":")
                figures.append((num.strip(), cap.strip()))
                i += 1; continue
            if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[i+1].strip()):
                j = i - 1
                while j >= 0 and not lines[j].strip():
                    j -= 1
                para_lines = []
                while j >= 0 and lines[j].strip():
                    para_lines.append(lines[j].strip()); j -= 1
                if para_lines:
                    tm = TABLE_CAP_RE.match(" ".join(reversed(para_lines)))
                    if tm:
                        tables.append((tm.group(1), tm.group(2).strip()))
                while i < len(lines) and lines[i].strip().startswith("|"):
                    i += 1
                continue
            i += 1
    return figures, tables


def write_caption_list(doc, heading, num_header, entries):
    """Render the List of Figures / List of Tables as a real three-column Word
    table (number, caption, page -- left blank for the author to fill in once
    the document is paginated), not a run of standalone paragraphs."""
    add_heading(doc, heading, "Heading 1")
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = num_header, "Caption", "Page"
    for c in hdr:
        for r in c.paragraphs[0].runs:
            r.bold = True
    for num, cap in entries:
        cells = t.add_row().cells
        cells[0].text = num
        cells[1].text = cap
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def centred(doc, text, size=None, bold=False, style="Normal"):
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    if size:
        r.font.size = Pt(size)
    return p


_bookmark_id = [1000]


def bookmark_run(run, name):
    """Wrap a run in a Word bookmark, so the header's REF LastName / REF
    PrimaryTitle fields (currently unresolved because clear_body() removed the
    template's own bookmarks along with its instructional text) have something
    to point back to again."""
    _bookmark_id[0] += 1
    bmid = str(_bookmark_id[0])
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bmid); start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bmid)
    run._r.addprevious(start)
    run._r.addnext(end)


def title_page(doc):
    for _ in range(3):
        doc.add_paragraph()
    centred(doc, TITLE, 20, True)
    # the header's REF PrimaryTitle field pulls from this bookmark, not the
    # printed title above -- kept short and hidden so the header fits on one
    # line even if the user refreshes fields (Ctrl+A, F9) in Word later
    hidden_p = doc.add_paragraph()
    hidden_r = hidden_p.add_run(SHORT_TITLE)
    hidden_r.font.hidden = True
    bookmark_run(hidden_r, "PrimaryTitle")
    centred(doc, SUBTITLE, 13)
    doc.add_paragraph()
    centred(doc, "submitted as a requirement for the degree of")
    centred(doc, PROGRAMME, bold=True)
    doc.add_paragraph()
    centred(doc, "Berlin, Germany")
    centred(doc, "by")
    pa = doc.add_paragraph(style="Normal")
    pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    first, _, last = AUTHOR.rpartition(" ")
    r1 = pa.add_run(first + " "); r1.bold = True; r1.font.size = Pt(13)
    r2 = pa.add_run(last); r2.bold = True; r2.font.size = Pt(13)
    bookmark_run(r2, "LastName")
    centred(doc, f"Matriculation Number {MATRIC}")
    doc.add_paragraph()
    centred(doc, f"Primary Thesis Supervisor: {SUPERVISOR}")
    centred(doc, f"Associate Thesis Supervisor: {ASSOCIATE}")
    doc.add_paragraph()
    centred(doc, "written in conjunction with Startup Insider GmbH")
    doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)


def fix_header_fields(doc):
    """The header's REF LastName / REF PrimaryTitle fields show their last
    cached result until Word recalculates them (Ctrl+A, F9) -- update that
    cached text directly too, so the header reads correctly even before that
    manual refresh, matching the bookmarks title_page() now provides."""
    for sec in doc.sections:
        for p in sec.header.paragraphs:
            for r in p.runs:
                if r.text == "<Last Name>":
                    r.text = AUTHOR.rpartition(" ")[-1]
                elif r.text == "<Primary Thesis Title>":
                    r.text = SHORT_TITLE


def main():
    doc = Document(str(TEMPLATE))
    doc.styles["Normal"].paragraph_format.line_spacing = 1.0  # template default was 1.5
    fix_header_fields(doc)
    clear_body(doc)
    title_page(doc)

    dictated_abstract = DICTATED / "00_abstract.md"
    if dictated_abstract.exists():
        text = dictated_abstract.read_text()
        text = re.sub(r"^## 0\.1 Abstract\n+", "", text, flags=re.M)  # drop the card subheading
        write_markdown(doc, text)
        doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)
    elif (ROOT / "abstract.md").exists():
        write_markdown(doc, (ROOT / "abstract.md").read_text())
        doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)

    used, chapter_texts = [], []
    for name in CHAPTERS:
        if (DICTATED / f"{name}.md").exists():
            src, tag = DICTATED / f"{name}.md", "dictated"
        elif (V2 / f"{name}.md").exists():
            src, tag = V2 / f"{name}.md", "v2"
        else:
            src, tag = V1 / f"{name}.md", "v1"
        used.append((name, tag))
        chapter_texts.append(src.read_text())

    appx = ROOT / "appendices.md"
    appx_text = appx.read_text() if appx.exists() else ""

    figures, tables = collect_captions(chapter_texts + [appx_text])
    write_caption_list(doc, "LIST OF FIGURES", "No.", figures)
    write_caption_list(doc, "LIST OF TABLES", "No.", tables)

    for text in chapter_texts:
        write_markdown(doc, text)
        doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)

    if appx_text:
        write_markdown(doc, appx_text)
        doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)

    refs_alpha = ROOT / "refs" / "references_alphabetical.md"
    if refs_alpha.exists():
        add_heading(doc, "Alphabetical List of References", "Heading 1")
        for line in refs_alpha.read_text().split("\n"):
            line = line.strip()
            if not line or line.startswith("#"):
                continue
            add_runs(doc.add_paragraph(style="Bibliography"), line)

    doc.save(str(OUT))
    words = sum(len(p.text.split()) for p in doc.paragraphs)
    print(f"{OUT.name}  —  {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables, ~{words:,} words")
    for n, v in used:
        print(f"   {n:<18} {v}")


if __name__ == "__main__":
    main()
